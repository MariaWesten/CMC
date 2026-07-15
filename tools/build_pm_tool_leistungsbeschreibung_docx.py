from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("outputs/Leistungsbeschreibung_PM_Tool_CMC.docx")

BLUE = "173F49"
BLUE_DARK = "0B2545"
MUTED = "52606D"
LIGHT_TEAL = "EDF7F4"
HEADER = "E8EEF5"
LIGHT_GRAY = "F6F8FA"
BORDER = "D9E2EC"
GREEN = "D9EAD3"
YELLOW = "FFF2CC"
RED = "F4CCCC"
BLUE_FILL = "D9EAF7"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color="1F2933", size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_col_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def repeat_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_para(doc, text="", style=None, bold=False, color=None, size=None, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        r.bold = bold
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
        if size:
            r.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.10
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.10
        p.add_run(item)


def add_note_box(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_borders(table, color="B7D5CF", size="6")
    set_cell_margins(table, top=120, bottom=120, start=160, end=160)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_TEAL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE_DARK)
    r.font.size = Pt(10)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    return p


def add_simple_table(doc, headers, rows, widths, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    set_col_widths(table, widths)
    for idx, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], HEADER)
        set_cell_text(table.rows[0].cells[idx], h, bold=True, color=BLUE_DARK, size=font_size)
    repeat_header_row(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), size=font_size)
            if value in ("Stark", "Sehr stark"):
                set_cell_shading(cells[idx], GREEN)
            elif value == "Erfüllt":
                set_cell_shading(cells[idx], BLUE_FILL)
            elif value in ("Teilweise", "Zu prüfen"):
                set_cell_shading(cells[idx], YELLOW)
            elif value in ("Schwach", "Risiko"):
                set_cell_shading(cells[idx], RED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    header = section.header.paragraphs[0]
    header.text = "CMC PM-Tool | Leistungsbeschreibung und Tool-Matrix"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Arbeitsstand 08.07.2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(8.5)
    footer.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    return doc


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = setup_doc()

    # Masthead
    add_para(doc, "LEISTUNGSBESCHREIBUNG", bold=True, color=BLUE, size=10, after=3)
    p = add_para(doc, "PM-Tool CMC", bold=True, color="000000", size=24, after=3)
    p.paragraph_format.line_spacing = 1.0
    add_para(
        doc,
        "Operative Steuerung statt Monday-Kopie: Anforderungen, Systemabgrenzung und Bewertungsmatrix für Asana, ClickUp und Zoho Projects",
        color=MUTED,
        size=12,
        after=14,
    )

    meta = [
        ("Stand", "08.07.2026"),
        ("Status", "Arbeitsversion für Toolauswahl und interne Abstimmung"),
        ("Kontext", "Monday soll abgelöst werden; neue Studie soll ab Anfang/Mitte August im neuen Tool umgesetzt werden."),
        ("Bestehende Systeme", "askDANTE für Zeiterfassung, Zoho als CRM, Microsoft 365 für Teams/SharePoint/Outlook."),
    ]
    table = add_simple_table(doc, ["Feld", "Angabe"], meta, [Inches(1.35), Inches(5.0)], font_size=9.5)
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True

    add_note_box(
        doc,
        "Kernaussage",
        "Das neue Tool soll nicht Monday 1:1 nachbauen. Es soll die operativen Fragen zuverlässig beantworten: Was ist heute zu tun, wie viel Zeit ist dafür geplant, wer ist aktuell verantwortlich, wo hängt die Aufgabe, wann wird sie fertig und an wen wird sie übergeben?",
    )

    heading(doc, "1. Zielbild", 1)
    add_para(
        doc,
        "Das PM-Tool ist die operative Steuerungsinstanz für Aufgaben, Durchlaufzeiten, Übergaben, Zeitbudgets und Status über mehrere CMC-Produktlinien hinweg. Es ersetzt keine Buchhaltung, kein CRM und kein Spezialreporting für NPS oder kaufmännische Kennzahlen.",
    )
    add_para(
        doc,
        "Die Auswahl sollte daher nicht anhand einer bestehenden Monday-Struktur erfolgen, sondern anhand der Frage, ob das Tool im Tagesgeschäft verlässlich steuerungsrelevante Antworten liefert.",
    )

    heading(doc, "2. Produktlinien und Scope", 1)
    add_para(doc, "Im Zielbild sollen mindestens folgende Produktlinien bzw. wiederkehrende Arbeitsarten abgebildet werden:")
    add_bullets(
        doc,
        [
            "PM Automation",
            "Creative Design",
            "CRM Audits",
            "Tagesgeschäft Mailing",
            "Studie",
            "Paketbeilagen",
            "optional: OKRs / Ziele",
        ],
    )

    heading(doc, "3. Abgrenzung zur bestehenden Systemlandschaft", 1)
    system_rows = [
        (
            "askDANTE",
            "Aktuell führend für Arbeits-/Projektzeiterfassung.",
            "Das PM-Tool muss Projektzeiten entweder selbst belastbar erfassen oder zumindest über Export/API/Mapping anschlussfähig sein. Da askDANTE perspektivisch abgelöst werden kann, ist native Aufgabenzeit ein wichtiges Auswahlkriterium.",
        ),
        (
            "Zoho CRM",
            "Führend für CRM-nahe Kunden-, Vertriebs- und ggf. kaufmännische Informationen.",
            "Das PM-Tool soll Kunden/Partner/Deals referenzieren können, aber keine CRM-Daten doppelt pflegen. Zoho Projects hat hier den strukturell stärksten Fit.",
        ),
        (
            "Microsoft 365",
            "Teams, SharePoint, Outlook und Dateiablage sind operativer Standard.",
            "Das PM-Tool sollte Teams-/Outlook-Benachrichtigungen und SharePoint/OneDrive-Dateilinks unterstützen. Die Datei-Wahrheit bleibt idealerweise in SharePoint.",
        ),
    ]
    add_simple_table(
        doc,
        ["System", "Heutige Rolle", "Anforderung an PM-Tool"],
        system_rows,
        [Inches(1.25), Inches(2.05), Inches(3.05)],
        font_size=8.8,
    )

    heading(doc, "4. Operative Kernfragen", 1)
    add_para(doc, "Das Tool muss ohne manuelle Nebenliste mindestens diese Fragen beantworten:")
    add_numbered(
        doc,
        [
            "Was sind heute meine Aufgaben und wie viel Zeit habe ich dafür?",
            "Bin ich mit der Aufgabe im Zeitbudget?",
            "Wo befindet sich die Aufgabe aktuell, also bei wem und in welchem Status?",
            "Wann ist die Aufgabe fertig und an wen wird sie anschließend übergeben?",
            "Wurde die Aufgabe in der vorgegebenen Zeit erledigt?",
            "Welche Aufgaben sind überfällig, blockiert oder budgetkritisch?",
        ],
    )

    heading(doc, "5. Muss-Anforderungen", 1)
    must_rows = [
        ("Aufgabenmodell", "Aufgaben, Unteraufgaben, Verantwortliche, Fälligkeitsdaten, Priorität, Status, Produktlinie, Kunde/Partner/Kampagne."),
        ("Zeitbudget", "Planzeit je Aufgabe bzw. Arbeitspaket; Ist-Zeit je Aufgabe; Soll-Ist-Abweichung auswertbar."),
        ("Status und Übergaben", "Aktueller Bearbeiter, nächster Bearbeiter, Übergabezeitpunkt, Blocker und Freigabestatus müssen klar sichtbar sein."),
        ("Persönliche Tagesansicht", "Jede Person sieht Aufgaben für heute inklusive geplanter Zeit und Priorität."),
        ("Workflow-Vorlagen", "Wiederverwendbare Templates je Produktlinie, z. B. Studie, Creative Design, Mailing, CRM Audit."),
        ("Dashboards", "Überfällige Aufgaben, Budgetrisiken, Status je Produktlinie, offene Übergaben und Auslastung müssen reportbar sein."),
        ("Integration/Export", "Kompatibilität mit Microsoft 365; Export/API für Reporting, Migration und mögliche Ablösung von askDANTE."),
        ("Rechte/Rollen", "Interne Rollen und ggf. externe Gäste/Kunden müssen steuerbar sein."),
    ]
    add_simple_table(doc, ["Bereich", "Mindestleistung"], must_rows, [Inches(1.7), Inches(4.65)], font_size=9)

    heading(doc, "6. Soll- und Kann-Anforderungen", 1)
    add_bullets(
        doc,
        [
            "Automatische Benachrichtigungen bei Statuswechsel, Blocker, Freigabe oder Übergabe.",
            "Formulare für neue Anfragen, Briefings und Produktlinien-spezifische Intake-Prozesse.",
            "Kapazitätsansicht über Personen und Teams mit geplanter Arbeitslast.",
            "Projekt-/Portfolioansicht über mehrere Produktlinien hinweg.",
            "OKR-/Zielmodul zur Verbindung von Zielen, Initiativen und operativen Aufgaben.",
            "Aufgabenbezogene Kommunikation mit Entscheidungs- und Freigabehistorie.",
            "Möglichkeit, SharePoint/OneDrive-Dateien zu verlinken statt Dateien redundant im PM-Tool zu pflegen.",
        ],
    )

    heading(doc, "7. Mindest-Datenmodell", 1)
    data_rows = [
        ("Identifikation", "Aufgaben-ID, Produktlinie, Projekt/Kunde/Partner, Kampagne/Studie, Arbeitspaket."),
        ("Steuerung", "Status, Priorität, Owner, aktueller Bearbeiter, nächster Bearbeiter, Übergabedatum."),
        ("Zeit", "Planzeit, Ist-Zeit, Restaufwand, Soll-Ist-Abweichung, Zeitbudget-Status."),
        ("Termine", "Startdatum, Fälligkeitsdatum, Meilenstein, Abhängigkeit, Eskalationsdatum."),
        ("Qualität", "Briefing vollständig, Assetcheck, QA, Freigabe, Fehler/Blocker, Abnahme."),
        ("Verweise", "Zoho-CRM-Link, SharePoint-Ordner, relevante Dateien, externe Partnerlinks."),
        ("Reporting", "Produktlinie, Team, Kunde/Partner, Monat/KW, Budgetstatus, Abschlussdatum."),
    ]
    add_simple_table(doc, ["Feldgruppe", "Beispiele"], data_rows, [Inches(1.45), Inches(4.9)], font_size=9)

    heading(doc, "8. Empfohlene Workflow-Vorlagen", 1)
    workflow_rows = [
        ("Studie", "Partnerbriefing -> Daten/Assets -> Variantenproduktion -> QA -> Freigabe -> Produktion -> Auswertung"),
        ("Tagesgeschäft Mailing", "Anfrage -> Briefing -> Umsetzung -> QA -> Versand-/Produktionsübergabe -> Reporting"),
        ("Creative Design", "Briefing -> Assetcheck -> Konzept -> Design -> Korrektur/Freigabe -> Export/Übergabe"),
        ("CRM Audits", "Anfrage -> Zugang/Datenbasis -> Analyse -> Findings -> Review -> Maßnahmenübergabe"),
        ("PM Automation", "Use Case -> Scoping -> Umsetzung -> Test -> Dokumentation -> Übergabe/Betrieb"),
        ("Paketbeilagen", "Anfrage -> Konzept -> Design -> Freigabe -> Produktion -> Versand-/Logistikabgleich"),
    ]
    add_simple_table(doc, ["Produktlinie", "Standard-Workflow"], workflow_rows, [Inches(1.55), Inches(4.8)], font_size=9)

    heading(doc, "9. Kommunikation im Tool", 1)
    add_para(
        doc,
        "Kommunikation im Tool ist sinnvoll, aber nicht zwingend als vollständiger Chat-Ersatz. Entscheidend ist: Entscheidungen, Freigaben, Blocker, Übergaben und finale Rückmeldungen müssen an der Aufgabe dokumentiert sein. Laufende Abstimmung kann weiter in Teams/E-Mail stattfinden, solange die operative Wahrheit im PM-Tool nachgezogen wird.",
    )

    heading(doc, "10. OKRs / Ziele", 1)
    add_para(
        doc,
        "OKRs sollten als optionale zweite Ausbaustufe betrachtet werden. Für den Start der neuen Studie ist wichtiger, dass Aufgaben, Zeitbudgets, Status und Übergaben funktionieren. Wenn OKRs genutzt werden, sollten sie oberhalb der Arbeit liegen: Ziel -> Initiative/Projekt -> Aufgaben.",
    )

    heading(doc, "11. Kosten- und Lizenzlogik", 1)
    add_para(
        doc,
        "Die reinen Lizenzpreise entscheiden den Fall nicht allein. Entscheidend ist, ab welcher Planstufe die CMC-Muss-Anforderungen ohne Workarounds abgebildet werden können. Sonst wirkt ein Tool günstig, erzeugt aber Kosten durch doppelte Pflege, fehlende Zeitauswertung oder parallele Nutzung von askDANTE.",
    )
    cost_rows = [
        (
            "Asana",
            "Starter: 10,99 USD/User/Monat jährlich; Advanced: 24,99 USD/User/Monat jährlich.",
            "Für CMC realistisch eher Advanced, weil Time Tracking, Goals, Workload, Portfolios/Formulas und stärkere Steuerung dort relevanter werden.",
            "Mittleres bis hohes Lizenzniveau, aber wahrscheinlich beste Team-Akzeptanz und geringerer Einführungswiderstand.",
        ),
        (
            "ClickUp",
            "Unlimited: 7 USD/User/Monat jährlich; Business: 12 USD/User/Monat jährlich.",
            "Für CMC realistisch eher Business, weil Portfolio Workload Management, Advanced Dashboards, stärkere Automationen und Proofing dort besser passen.",
            "Gutes Preis-Leistungs-Verhältnis, aber höheres Risiko durch Setup-Komplexität und Governance-Aufwand.",
        ),
        (
            "Zoho Projects",
            "Premium: 4 EUR/User/Monat jährlich; Enterprise: 9 EUR/User/Monat jährlich; Ultimate: 14 EUR/User/Monat jährlich.",
            "Für CMC mindestens Premium; Enterprise prüfen, wenn Portfolio-Dashboard, Rollen/Felder, SSO, kritischer Pfad/Baseline und stärkere Governance wichtig sind.",
            "Preislich sehr attraktiv und systemisch passend zu Zoho CRM; Akzeptanz und Bedienbarkeit im Team sind die Hauptprüfung.",
        ),
    ]
    add_simple_table(
        doc,
        ["Tool", "Preisanker", "Voraussichtlich relevanter Plan", "Kosten-Einordnung"],
        cost_rows,
        [Inches(1.15), Inches(1.75), Inches(2.0), Inches(1.45)],
        font_size=8.2,
    )
    add_note_box(
        doc,
        "Kostenfazit",
        "Zoho Projects ist auf Lizenzebene voraussichtlich am günstigsten, ClickUp bietet wahrscheinlich den besten Funktionsumfang je Euro/Dollar, Asana ist vermutlich am teuersten für den vollen CMC-Fit, kann aber durch bessere Akzeptanz und geringeren Schulungsaufwand trotzdem wirtschaftlich sein.",
    )

    # Landscape matrix section
    matrix_section = doc.add_section(WD_SECTION.NEW_PAGE)
    matrix_section.orientation = WD_ORIENT.LANDSCAPE
    matrix_section.page_width, matrix_section.page_height = matrix_section.page_height, matrix_section.page_width
    matrix_section.top_margin = Inches(0.65)
    matrix_section.bottom_margin = Inches(0.65)
    matrix_section.left_margin = Inches(0.65)
    matrix_section.right_margin = Inches(0.65)

    heading(doc, "12. Bewertungsmatrix der drei Favoriten", 1)
    add_para(
        doc,
        "Legende: Stark = sehr guter Fit für CMC; Erfüllt = grundsätzlich passend; Teilweise = lösbar, aber mit Einschränkung oder Konfiguration; Zu prüfen = Anbieter-/Setup-Frage vor Entscheidung klären.",
        color=MUTED,
        size=9.5,
        after=6,
    )

    matrix_rows = [
        ("Tagesansicht: Was muss ich heute tun?", "Erfüllt", "Stark", "Erfüllt", "Alle drei geeignet. ClickUp ist sehr stark bei persönlichem Arbeitsbereich; Asana ist klar und adoption-freundlich; Zoho solide."),
        ("Planzeit je Aufgabe", "Erfüllt", "Stark", "Stark", "ClickUp und Zoho sind für Schätzungen/Zeitbudgets besonders naheliegend; Asana kann es, aber ggf. weniger tief im Standardreporting."),
        ("Ist-Zeit je Aufgabe", "Erfüllt", "Stark", "Stark", "Wichtig wegen askDANTE-Ablöse. ClickUp/Zoho wirken stärker als mögliche spätere Projektzeiterfassung."),
        ("Soll-Ist-Auswertung Zeitbudget", "Teilweise", "Stark", "Stark", "Zoho bietet geplante vs. tatsächliche Stunden/Budgetberichte; ClickUp mit Time Estimates/Time Tracking gut; Asana prüfen im Zielplan."),
        ("Status, Owner, aktueller Bearbeiter", "Stark", "Stark", "Erfüllt", "Asana und ClickUp sehr klar für operative Statusarbeit; Zoho erfüllt, wirkt aber stärker projekt-/zeitwirtschaftlich."),
        ("Übergabe an nächste Person", "Stark", "Stark", "Erfüllt", "Über Custom Fields, Automationen und Statuswechsel abbildbar. In Zoho ebenfalls möglich, aber Setup prüfen."),
        ("Abhängigkeiten, Meilensteine, Gantt", "Stark", "Stark", "Stark", "Alle drei geeignet für Studienplanung und wiederkehrende Projektlogiken."),
        ("Vorlagen je Produktlinie", "Stark", "Stark", "Stark", "Alle drei können wiederkehrende Workflows standardisieren."),
        ("Automationen bei Statuswechsel", "Stark", "Stark", "Stark", "Alle drei können Übergaben/Benachrichtigungen automatisieren; Detailtiefe im Pilot prüfen."),
        ("Formulare / Intake", "Stark", "Stark", "Erfüllt", "Asana/ClickUp stark für strukturierte Anfragen; Zoho auch möglich, besonders im Zoho-Umfeld interessant."),
        ("Dashboards / Portfolio", "Stark", "Stark", "Stark", "Alle drei geeignet. Asana stark bei Portfolio/Managementsicht, ClickUp flexibel, Zoho mit Budget-/Projektstatus nah an CRM/Finanzen."),
        ("Kapazitäts-/Workloadplanung", "Stark", "Stark", "Erfüllt", "Asana/ClickUp stark für Teamlast; Zoho ebenfalls vorhanden, Detailtiefe und Bedienbarkeit testen."),
        ("OKRs / Ziele", "Stark", "Stark", "Teilweise", "Asana und ClickUp sind für Zielverknüpfung geeigneter. Zoho eher über Projekte/Reports statt OKR-nativ."),
        ("Kommunikation an Aufgaben", "Erfüllt", "Stark", "Erfüllt", "ClickUp bietet zusätzlich Chat; bei Asana/Zoho reicht aufgabenbezogene Kommentierung für operative Wahrheit."),
        ("Microsoft 365 / Teams / SharePoint", "Erfüllt", "Erfüllt", "Erfüllt", "Alle grundsätzlich kompatibel. Zoho und ClickUp nennen Microsoft-Integrationen breit; Asana hat großes App-Ökosystem. SharePoint-Dateiwahrheit im Pilot testen."),
        ("Zoho CRM Anschluss", "Teilweise", "Teilweise", "Stark", "Zoho Projects hat klaren Systemvorteil, wenn Zoho CRM langfristig gesetzt ist."),
        ("askDANTE Übergang / spätere Ablöse", "Teilweise", "Stark", "Stark", "askDANTE hat Projektzeiterfassung und REST-API; für Ablöse sind ClickUp/Zoho als Zielsystem naheliegender als Asana."),
        ("Lizenzkosten für vollen CMC-Fit", "Schwach", "Stark", "Sehr stark", "Zoho ist am günstigsten; ClickUp liegt funktional/preislich stark; Asana wird für Zeit/Workload/Goals wahrscheinlich erst im Advanced-Plan wirklich rund."),
        ("Kostenrisiko Einführung", "Stark", "Teilweise", "Erfüllt", "Asana dürfte wegen Team-Akzeptanz am wenigsten Reibung erzeugen. ClickUp kann viel, braucht aber strikte Konfiguration. Zoho braucht Akzeptanztest im Team."),
        ("Kostenrisiko Parallelbetrieb askDANTE", "Teilweise", "Stark", "Stark", "Wenn Zeiterfassung mittelfristig aus askDANTE heraus soll, sind ClickUp/Zoho wirtschaftlich interessanter. Bei Asana kann askDANTE länger nötig bleiben."),
        ("Einführungsaufwand", "Stark", "Teilweise", "Erfüllt", "Asana vermutlich am leichtesten einzuführen; ClickUp braucht Governance wegen Funktionsfülle; Zoho braucht Prozess-/CRM-Konzept."),
        ("Gesamtfit CMC", "Stark", "Stark", "Stark", "Kein Tool fällt fachlich raus. Entscheidung hängt an Priorität: Adoption/OKRs = Asana; All-in-one/Zeit = ClickUp; Zoho-Stack/Zeitbudget = Zoho Projects."),
    ]
    add_simple_table(
        doc,
        ["Anforderung", "Asana", "ClickUp", "Zoho Projects", "CMC-Einordnung"],
        matrix_rows,
        [Inches(2.15), Inches(0.92), Inches(0.92), Inches(1.05), Inches(4.55)],
        font_size=7.7,
    )

    heading(doc, "13. Entscheidungsempfehlung", 1)
    add_para(
        doc,
        "Die Entscheidung muss nicht bedeuten, dass sich Maria in drei Tools tief einarbeitet. Sinnvoller ist eine Vorentscheidung anhand von Knock-out-Fragen und danach ein sehr kleiner Pilot mit einem Favoriten und einem Backup.",
    )
    recommendation_rows = [
        ("Asana", "Empfehlen, wenn Team-Akzeptanz und schnelle Einführung das wichtigste Risiko sind. Wirtschaftlich nur dann stark, wenn der höhere Advanced-Preis durch weniger Reibung und klare Nutzung gerechtfertigt ist."),
        ("ClickUp", "Empfehlen, wenn CMC ein PM-/Zeiterfassungs-/Workflow-System aus einem Guss sucht und bereit ist, das Setup eng zu führen. Funktional vermutlich stärkster Match für die Anforderungen."),
        ("Zoho Projects", "Empfehlen, wenn Zoho als bestehender Stack strategisch genutzt werden soll und Lizenzkosten/Zeiterfassung stark zählen. Vorher Bedienbarkeit und Team-Akzeptanz testen."),
    ]
    add_simple_table(doc, ["Tool", "Dann bevorzugen"], recommendation_rows, [Inches(1.35), Inches(8.25)], font_size=8.6)
    decision_rows = [
        (
            "Wenn Robert/Team sagt: Akzeptanz schlägt alles",
            "Asana als Empfehlung, ClickUp nur als funktional stärkere Alternative erwähnen.",
        ),
        (
            "Wenn Zeiterfassung und askDANTE-Ablöse zentral sind",
            "ClickUp oder Zoho Projects priorisieren; Asana nur prüfen, wenn askDANTE vorerst bleibt.",
        ),
        (
            "Wenn Zoho ohnehin ausgebaut werden soll",
            "Zoho Projects als wirtschaftlich-systemische Empfehlung; Pilot muss Bedienbarkeit beweisen.",
        ),
        (
            "Wenn 'kann am meisten' das wichtigste Argument ist",
            "ClickUp empfehlen, aber mit klarer Governance: wenige Views, feste Felder, keine Tool-Spielwiese.",
        ),
    ]
    add_simple_table(
        doc,
        ["Entscheidungslogik", "Empfehlung"],
        decision_rows,
        [Inches(3.0), Inches(6.6)],
        font_size=8.8,
    )

    heading(doc, "14. Pilotvorschlag bis Anfang/Mitte August", 1)
    add_numbered(
        doc,
        [
            "Nicht drei Tools vollständig einarbeiten: zuerst 30-Minuten-Demos nach denselben sechs operativen Kernfragen durchführen.",
            "Danach einen Hauptfavoriten und ein Backup wählen.",
            "Im Hauptfavoriten 10-15 reale Aufgaben aus Studie und Tagesgeschäft Mailing anlegen: Planzeit, Owner, Status, Deadline, Übergabe, Datei-/Zoho-Link.",
            "Eine Woche lang Tagesansicht, Zeiterfassung, Übergaben und Statusreporting testen.",
            "Nur wenn der Hauptfavorit die sechs Kernfragen nicht sauber beantwortet, Backup testen.",
        ],
    )

    heading(doc, "15. Quellen und Prüfstand", 1)
    source_intro = add_para(
        doc,
        "Öffentliche Herstellerinformationen, geprüft am 08.07.2026. Feature- und Planverfügbarkeit sollten vor Vertragsentscheidung direkt beim Anbieter bzw. in der konkreten Lizenzstufe verifiziert werden:",
        size=9.5,
        color=MUTED,
    )
    sources = [
        ("Asana Project Management", "https://asana.com/features/project-management"),
        ("Asana Goals & Reporting", "https://asana.com/features/goals-reporting"),
        ("Asana Pricing", "https://asana.com/pricing"),
        ("ClickUp Features", "https://clickup.com/features"),
        ("ClickUp Pricing", "https://clickup.com/pricing"),
        ("Zoho Projects Features", "https://www.zoho.com/de/projects/features.html"),
        ("Zoho Projects Pricing", "https://www.zoho.com/de/projects/zohoprojects-pricing.html"),
        ("askDANTE Zeiterfassung", "https://www.askdante.com/"),
        ("Microsoft Planner", "https://www.microsoft.com/en-us/microsoft-365/planner/microsoft-planner"),
    ]
    for label, url in sources:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(label + ": ")
        add_hyperlink(p, url, url)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
