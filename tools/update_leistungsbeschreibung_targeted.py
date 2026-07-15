from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document


PATH = Path("outputs/Leistungsbeschreibung_URL_QR_Code_Tool.docx")
BACKUP = Path("outputs/Leistungsbeschreibung_URL_QR_Code_Tool.backup-before-targeted-edit.docx")


def set_cell_text(cell, text):
    cell.text = ""
    p = cell.paragraphs[0]
    p.add_run(text)


def set_row(row, values):
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)


def insert_row_after(table, after_label, values):
    for row in table.rows:
        if row.cells[0].text.strip() == after_label:
            new_tr = deepcopy(row._tr)
            row._tr.addnext(new_tr)
            new_row = table.rows[list(table._tbl.tr_lst).index(new_tr)]
            set_row(new_row, values)
            return True
    return False


def replace_paragraph(doc, old_start, new_text):
    for p in doc.paragraphs:
        if p.text.strip().startswith(old_start):
            style = p.style
            p.clear()
            p.style = style
            p.add_run(new_text)
            return True
    return False


def replace_exact_paragraph(doc, old_text, new_text):
    for p in doc.paragraphs:
        if p.text.strip() == old_text:
            style = p.style
            p.clear()
            p.style = style
            p.add_run(new_text)
            return True
    return False


def main():
    if not BACKUP.exists():
        shutil.copy2(PATH, BACKUP)

    doc = Document(PATH)

    replace_paragraph(
        doc,
        "Das URL-/QR-Code-Tool soll Malin langfristig",
        "Das URL-/QR-Code-Tool soll Malin langfristig als zentrale Lösung für die kampagnenweise Erstellung, Verwaltung und Auswertung individueller Tracking-Links ablösen. Der Schwerpunkt liegt nicht auf einem klassischen QR-Code-Design-Tool, sondern auf stabilen, massenhaft erzeugbaren Tracking-Links. Zusätzlich soll das Tool die Trackingqualität erhöhen, indem Links konsistent erzeugt, vorab validiert und später sauber pro Kampagne ausgewertet werden können.",
    )
    replace_exact_paragraph(
        doc,
        "Das Tool soll den bestehenden CSM-/Data-/Produktionsprozess zuverlässig abbilden und gleichzeitig genug Struktur schaffen, um Fehler vor dem Versand von Mailings zu vermeiden.",
        "Das Tool soll den bestehenden CSM-/Data-/Produktionsprozess zuverlässig abbilden und gleichzeitig genug Struktur schaffen, um Fehler vor dem Versand von Mailings zu vermeiden. Entscheidend ist dabei, dass gedruckte oder bereits versendete QR-/Tracking-Links stabil bleiben.",
    )

    # Funktionsumfang table
    table = doc.tables[0]
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label == "Redirect":
            set_row(
                row,
                [
                    "Redirect",
                    "Zuverlässige Weiterleitung vom Tracking-Link zur hinterlegten Ziel-URL ohne Werbung, Delay oder externe Zwischenseite. Zu klären ist, ob Redirects für eine definierte Dauer oder dauerhaft bzw. so lange wie fachlich benötigt aufrechterhalten werden sollen.",
                    "Muss-Funktion, da fehlerhafte oder ablaufende Links Mailings und gedruckte QR-Codes direkt gefährden.",
                ],
            )
        elif label == "QR-Code-Ausgabe":
            set_row(
                row,
                [
                    "QR-Code-Ausgabe",
                    "Es muss nicht zwingend je Link eine statische Bilddatei erzeugt werden. Ausreichend wäre auch ein Bildlink bzw. ein QR-Code-Endpunkt, bei dem das QR-Bild bei Abruf generiert wird. Ein ZIP-/Bildexport kann optional bleiben.",
                    "Wichtig für Produktion, aber bei 10.000-100.000 Links pro Kampagne wahrscheinlich besser über Bildlink/on-demand als über massenhaften Bilddatei-Export.",
                ],
            )
        elif label == "Reporting":
            set_row(
                row,
                [
                    "Reporting",
                    "Klick-/Scan-Auswertung pro Kampagne, Link, Code und Zeitraum; zusätzlich exportierbare Roh- oder Aggregatdaten. Die Datenstruktur soll die Qualität und Nachvollziehbarkeit des Trackings verbessern.",
                    "Muss-Funktion, Detailtiefe mit Data/CSM zu klären.",
                ],
            )

    if not any(r.cells[0].text.strip() == "Ziel-URL nach Druck änderbar" for r in table.rows):
        insert_row_after(
            table,
            "Redirect",
            [
                "Ziel-URL nach Druck änderbar",
                "Zu prüfen ist, ob die Ziel-URL eines Tracking-Links nachträglich geändert werden kann, ohne dass sich der bereits gedruckte oder versendete QR-Code/Tracking-Link ändert.",
                "Wichtiger Vorteil dynamischer Links; besonders relevant, wenn Fehler nach Druck/Versand auffallen oder Zielseiten später wechseln.",
            ],
        )

    # Mengenmodell table
    table = doc.tables[1]
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label == "Link-Laufzeit":
            set_row(
                row,
                [
                    "Link-Laufzeit",
                    "Zu klären: definierte Dauer, dauerhaft oder so lange wie Kampagne/Printmaterial fachlich benötigt wird.",
                    "Relevant für Datenhaltung, Archivierung, Kosten und gedruckte QR-Codes.",
                ],
            )

    # Short.io question table
    table = doc.tables[2]
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label == "Redirects":
            set_row(
                row,
                [
                    "Redirects",
                    "Sind Redirects dauerhaft stabil, ohne Werbung, Delay oder Zwischenseite? Für welche Dauer bleiben Links aktiv: definierte Laufzeit, dauerhaft oder solange der Account besteht? Kann die Ziel-URL nach Druck/Versand geändert werden, ohne den QR-Code/Tracking-Link zu ändern?",
                ],
            )
        elif label == "QR-Codes":
            set_row(
                row,
                [
                    "QR-Codes",
                    "Müssen QR-Code-Bilder je Link als Datei erzeugt werden oder können Bildlinks/QR-Endpunkte genutzt werden, bei denen das QR-Bild bei Abruf generiert wird? Gibt es Limits für Abrufe, Downloads, Exporte oder QR-Code-Generierung bei 10.000-100.000 Links je Kampagne?",
                ],
            )
        elif label == "Reporting":
            set_row(
                row,
                [
                    "Reporting / Trackingqualität",
                    "Welche Klick-/Scan-Daten sind verfügbar? Können Reports pro Kampagne, Link und Zeitraum exportiert werden? Unterstützt Short.io eine saubere Kampagnenstruktur, konsistente Parameter und Exporte, sodass die Trackingqualität gegenüber dem heutigen Prozess verbessert wird? Gibt es Eventlimits oder Aufbewahrungsfristen?",
                ],
            )

    replace_exact_paragraph(
        doc,
        "8. Punkte, die intern vor der Anfrage finalisiert werden sollten",
        "8. Offene Punkte / Annahmen für die Anfrage",
    )
    replace_exact_paragraph(
        doc,
        "Finales CSV-Format: Pflichtfelder, optionale Felder, Code-/URL-Logik und erwartete Exportspalten.",
        "Finales CSV-Format: Pflichtfelder, optionale Felder, Code-/URL-Logik und erwartete Exportspalten. Das muss nicht zwingend vor der ersten Short.io-Anfrage final beantwortet sein, sollte aber als offener Klärpunkt mitgedacht werden.",
    )
    replace_exact_paragraph(
        doc,
        "Heutiges und mittelfristiges Mengengerüst: Kampagnen, Links, Klicks/Scans, Laufzeit der Links.",
        "Heutiges und mittelfristiges Mengengerüst: Kampagnen, Links, Klicks/Scans und Laufzeit der Links. Für die Anfrage sollte zunächst bestätigt werden, ob ca. 20 Kampagnen pro Monat und 10.000-100.000 Links pro Kampagne als Annahme verwendet werden sollen.",
    )
    replace_exact_paragraph(
        doc,
        "Reporting-Erwartung: Welche Auswertungen braucht CSM/Data wirklich?",
        "Reporting-Erwartung: Welche Auswertungen braucht CSM/Data wirklich? Das kann parallel zur Anbieterprüfung weiter geschärft werden.",
    )
    replace_exact_paragraph(
        doc,
        "Robert prüft, ob der oben beschriebene Funktionsumfang dem erwarteten Zielbild entspricht.",
        "Robert prüft, ob der oben beschriebene Funktionsumfang dem erwarteten Zielbild entspricht und ob noch Anforderungen oder Fragen fehlen, die in den Fragen an Short.io ergänzt werden sollen.",
    )
    replace_exact_paragraph(
        doc,
        "Nach Freigabe wird Short.io mit dieser Leistungsbeschreibung angefragt.",
        "Nach Rückmeldung zu Mengengerüst und offenen Anforderungen wird Short.io mit dieser Leistungsbeschreibung angefragt.",
    )

    doc.save(PATH)
    print(PATH)


if __name__ == "__main__":
    main()
