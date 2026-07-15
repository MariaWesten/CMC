from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "print_mailing_ki_produktion" / "CMC_KI_Prozess_Print_Mailing_Studie_2027.docx"

INK = "183F47"
ACCENT = "2F8F83"
LIGHT = "EAF4F2"
PALE = "F4F7F7"
MUTED = "65767C"
WHITE = "FFFFFF"
LINE = "D9E2E4"
WARN = "FFF4D6"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=11, bold=False, color="1F2D32", italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])
    set_font(run, size=9, color=MUTED)


def add_para(doc, text="", *, bold_prefix=None, style=None, after=6, keep=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if bold_prefix and text.startswith(bold_prefix):
        lead = p.add_run(bold_prefix)
        set_font(lead, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_font(rest)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_bullet(doc, text):
    p = add_para(doc, text, style="List Bullet", after=4)
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    return p


def add_number(doc, text):
    p = add_para(doc, text, style="List Number", after=5)
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    return p


def add_callout(doc, label, text, fill=LIGHT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    lead = p.add_run(label + " ")
    set_font(lead, bold=True, color=INK)
    body = p.add_run(text)
    set_font(body)
    return p


def add_step(doc, number, title, body, owner):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [920, 8440])
    left, right = table.rows[0].cells
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    set_cell_shading(left, ACCENT)
    set_cell_shading(right, PALE)
    set_cell_margins(left, 120, 100, 120, 100)
    set_cell_margins(right, 120, 170, 120, 170)

    p_num = left.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_num.paragraph_format.space_after = Pt(0)
    run = p_num.add_run(str(number))
    set_font(run, size=18, bold=True, color=WHITE)

    p_title = right.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(3)
    run = p_title.add_run(title)
    set_font(run, size=12, bold=True, color=INK)
    p_body = right.add_paragraph()
    p_body.paragraph_format.space_after = Pt(3)
    p_body.paragraph_format.line_spacing = 1.15
    set_font(p_body.add_run(body), size=10.5)
    p_owner = right.add_paragraph()
    p_owner.paragraph_format.space_after = Pt(0)
    set_font(p_owner.add_run("Verantwortung: "), size=9.5, bold=True, color=MUTED)
    set_font(p_owner.add_run(owner), size=9.5, color=MUTED)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after in (
        ("Heading 1", 16, 18, 8),
        ("Heading 2", 13, 14, 7),
        ("Heading 3", 11.5, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(INK if name == "Heading 1" else ACCENT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    set_font(hp.add_run("CMC Print-Mailing-Studie 2027  |  Prozessdokumentation"), size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(32)
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("PROZESSDOKUMENTATION"), size=10, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run("KI-gestützte Produktion der\nCMC Print-Mailing-Studie 2027"), size=27, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    set_font(p.add_run("Aktueller Prototyp und operativer Ablauf"), size=14, color=MUTED)

    add_callout(
        doc,
        "Zielbild:",
        "Vier partnerindividuelle Layoutvarianten mit Vorder- und Rückseite werden aus strukturierten Partnerdaten, KI-Texten und standardisierten InDesign-Templates erzeugt. Das Ergebnis bleibt vollständig in InDesign bearbeitbar.",
    )

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(meta, [2200, 7160])
    for row in meta.rows:
        set_cell_margins(row.cells[0], 90, 120, 90, 120)
        set_cell_margins(row.cells[1], 90, 120, 90, 120)
        set_cell_shading(row.cells[0], PALE)
    values = [
        ("Stand", "09.07.2026"),
        ("Status", "Funktionsfähiger Prototyp; InDesign-Test und weitere Zeitmessungen laufen"),
        ("Geltungsbereich", "Partnerinput bis befüllte InDesign-Datei; manuelles Layout-Finishing und Qualitätsprüfung bleiben Bestandteil des Prozesses"),
    ]
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].text = ""
        row.cells[1].text = ""
        set_font(row.cells[0].paragraphs[0].add_run(label), size=10, bold=True, color=INK)
        set_font(row.cells[1].paragraphs[0].add_run(value), size=10)

    doc.add_page_break()


def build_document():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("1. Prozess auf einen Blick", level=1)
    add_para(
        doc,
        "Der Prozess kombiniert standardisierte Automatisierung mit bewusst gesetzten manuellen Qualitätsstufen. Die KI erstellt keine fertige Gestaltung, sondern strukturierte, slotfähige Texte. Der Mailing-Assistent prüft und bündelt die Daten; das InDesign-Skript übernimmt die technische Befüllung.",
    )
    add_callout(
        doc,
        "Wichtig:",
        "Der aktuelle Prototyp erzeugt die Texte nicht direkt im Mailing-Assistenten. Er erstellt einen vollständigen KI-Textauftrag, der zusammen mit dem Partnerordner im freigegebenen KI-Chat ausgeführt wird. Das Ergebnis wird anschließend als JSON importiert.",
        fill=WARN,
    )

    steps = [
        ("Partnerinput bereitstellen", "Briefing, Angebot, CI, Bilder, Logos, Produkte und Pflichttexte werden als Partnerordner bzw. ZIP bereitgestellt.", "Partner / CSM / Operations"),
        ("KI-Textauftrag erzeugen", "Partner-ZIP laden, Website und optionale Textvorgaben eintragen und den standardisierten Textauftrag herunterladen.", "Jule oder Simone"),
        ("Texte mit KI erstellen", "Partner-ZIP und Textauftrag gemeinsam im freigegebenen KI-Chat hochladen. Die KI liefert ausschließlich das definierte Text-JSON.", "Jule oder Simone mit KI-Unterstützung"),
        ("Daten prüfen und zuordnen", "KI-JSON direkt einfügen oder importieren, Maximalzeichen prüfen, fehlende Angaben erkennen und Bilder den vorgesehenen Slots zuordnen.", "Jule oder Simone"),
        ("InDesign automatisiert befüllen", "Standardtemplate öffnen, Skript ausführen und die exportierte Partner-JSON wählen. Texte und Bilder werden in die benannten Rahmen gesetzt.", "Jule oder Simone"),
        ("Layout und Druckdaten finalisieren", "Bildauswahl, Ausschnitte, Typografie, Übersatz, Pflichttexte, Reinzeichnung, Peer Review und Andruck werden fachlich geprüft.", "Design / Peer Review / Operations"),
    ]
    for idx, (title, body, owner) in enumerate(steps, 1):
        add_step(doc, idx, title, body, owner)

    doc.add_heading("2. Benötigte Eingaben", level=1)
    add_bullet(doc, "Partnerbriefing mit Marke, Zielgruppe, Produktfokus, Angebot, Gültigkeit, Anrede und Pflichttexten")
    add_bullet(doc, "CI-Unterlagen, Logo, Schriften, Farben und freigegebene Gestaltungselemente")
    add_bullet(doc, "Produkt-, Mood- und Absenderbilder in ausreichender Druckqualität")
    add_bullet(doc, "Website-URL und bei Bedarf ein kurzer Referenztext für Tonalität und Wording")
    add_bullet(doc, "Echte Kundenstimmen mit Namen, sofern Zitatmodule verwendet werden")
    add_bullet(doc, "Standardisierte InDesign-Datei mit eindeutig benannten Text- und Bildrahmen")

    doc.add_heading("3. Texterzeugung im aktuellen Prototyp", level=1)
    add_para(
        doc,
        "Die Texterzeugung ist als kontrollierter Übergabeschritt organisiert. Dadurch können Jule und Simone die Texte selbst erzeugen, ohne JSON manuell aufzubauen und ohne Zugangsdaten in einem lokalen Tool zu speichern.",
    )
    add_number(doc, "Partner-ZIP im CMC Mailing-Assistenten laden.")
    add_number(doc, "Website-URL eintragen und bei Bedarf partnerspezifische Vorgaben ergänzen.")
    add_number(doc, "Verbindliche Botschaften, Claims und No-Gos eintragen.")
    add_number(doc, "Den KI-Textauftrag herunterladen.")
    add_number(doc, "Im freigegebenen KI-Chat sowohl den Partnerordner als auch den KI-Textauftrag hochladen.")
    add_number(doc, "Die vollständige JSON-Antwort kopieren und im Mailing-Assistenten einfügen; alternativ eine gespeicherte JSON-Datei importieren.")
    add_number(doc, "Texte, Fakten, Platzhalter und Maximalzeichen fachlich prüfen.")

    doc.add_heading("Steuerungsmöglichkeiten", level=2)
    add_bullet(doc, "Tonalität je Partner, zum Beispiel sachlicher, emotionaler, hochwertiger oder reduzierter")
    add_bullet(doc, "Verbindliche Markenbegriffe, Produktnamen und Kernbotschaften")
    add_bullet(doc, "Nicht zulässige Begriffe, Superlative oder Claims")
    add_bullet(doc, "Website-Wording als Referenz für Ansprache und Sprachrhythmus")
    add_bullet(doc, "Unterschiedliche Kommunikationslogik für A Klassischer Werbebrief, B Visual Storytelling, C Snackable und D Editorial")
    add_bullet(doc, "Harte feldbezogene Maximalzeichen inklusive Leerzeichen")

    add_callout(
        doc,
        "Regel für Zitate:",
        "Kundenzitate werden niemals erfunden. Sie dürfen nur aus dem Partnerbriefing oder eindeutig belegbar von der Partnerwebsite übernommen werden. Fehlen echte Zitate, markiert der Mailing-Assistent die Felder für eine Rückfrage.",
    )

    doc.add_heading("4. Aufgabe des CMC Mailing-Assistenten", level=1)
    add_bullet(doc, "Partner-ZIP analysieren und Bilddateien lokal für InDesign bereitstellen")
    add_bullet(doc, "Standardisierten KI-Textauftrag mit Variantenlogik und Zeichenlimits erzeugen")
    add_bullet(doc, "KI-Text-JSON direkt einfügen oder importieren, ohne Bildpfade oder technische Optionen zu überschreiben")
    add_bullet(doc, "Alle Textfelder editierbar anzeigen und Längenampeln berechnen")
    add_bullet(doc, "Fehlende Pflichtangaben, Platzhalter und doppelte Umbrüche markieren")
    add_bullet(doc, "Bilder automatisch vorschlagen und eine manuelle Korrektur ermöglichen")
    add_bullet(doc, "Partner-Rückfrage, Checkreport und finale InDesign-JSON exportieren")

    doc.add_heading("5. Aufgabe des InDesign-Skripts", level=1)
    add_para(
        doc,
        "Das Skript arbeitet mit den Namen der Text- und Bildrahmen. Die JSON-Schlüssel und die InDesign-Rahmennamen müssen daher exakt übereinstimmen.",
    )
    add_bullet(doc, "Texte in alle gleichnamigen Rahmen einsetzen")
    add_bullet(doc, "Platzhalter innerhalb von Template-Texten ersetzen")
    add_bullet(doc, "Leere JSON-Werte überspringen, damit vorhandene Template-Texte erhalten bleiben")
    add_bullet(doc, "Bilder proportional und mittig vollflächig in den Rahmen einsetzen")
    add_bullet(doc, "Eine weiterhin normal bearbeitbare InDesign-Kopie nach der definierten Dateinamenslogik speichern")
    add_bullet(doc, "Übersatz, fehlende Rahmen, fehlende Dateien und weitere Auffälligkeiten melden")

    doc.add_heading("6. Automatisiert und manuell", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [4680, 4680])
    headers = ("Automatisiert oder teilautomatisiert", "Weiterhin fachlich/manuell")
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, INK)
        set_cell_margins(cell)
        cell.text = ""
        set_font(cell.paragraphs[0].add_run(header), size=10, bold=True, color=WHITE)
    rows = [
        ("Prompt- und Feldlogik", "Vollständigkeit und Plausibilität des Briefings"),
        ("Textstruktur und Zeichenlimits", "Bewertung von Marke, Zielbild und Produktfokus"),
        ("Import und Zusammenführung der JSON", "Auswahl und gestalterische Beurteilung der Bilder"),
        ("Bildpfade und technische Rahmenbefüllung", "Layout-Finishing, Typografie und Bildausschnitte"),
        ("Pflichtfeld-, Platzhalter- und Übersatzhinweise", "Rechtliche und inhaltliche Prüfung"),
        ("Dateiname, Kopie und Befüllreport", "Peer Review, Reinzeichnung und Andruckprüfung"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell)
        set_font(cells[0].paragraphs[0].add_run(left), size=9.5)
        set_font(cells[1].paragraphs[0].add_run(right), size=9.5)

    doc.add_heading("7. Qualitätsgates", level=1)
    quality_items = (
        "Angebot, Incentive, Produktfokus und Pflichttexte sind über alle Varianten konsistent.",
        "Keine unbelegten Produktversprechen, Wirkversprechen oder Kundenzitate.",
        "Alle Pflichtfelder und Pflichtbilder sind vorhanden oder als Rückfrage dokumentiert.",
        "Kein Text überschreitet die Maximalzeichen aus der Rahmentabelle.",
        "Keine sichtbaren technischen Platzhalter oder doppelten Absatzabstände.",
        "Keine Übersatztexte; Bilder sind korrekt platziert und druckfähig.",
        "Peer Review, Reinzeichnung und Andruckprüfung sind abgeschlossen.",
    )
    for index, item in enumerate(quality_items):
        paragraph = add_bullet(doc, item)
        paragraph.paragraph_format.keep_with_next = index < len(quality_items) - 1

    doc.add_heading("8. Aktueller Stand und nächste Ausbaustufe", level=1)
    add_callout(
        doc,
        "Heute nutzbar:",
        "Partner-ZIP analysieren, KI-Textauftrag erzeugen, KI-Ergebnis importieren, Texte und Bilder prüfen, JSON exportieren und InDesign automatisiert befüllen.",
    )
    add_callout(
        doc,
        "Noch nicht vollautomatisch:",
        "Die KI-Ausführung erfolgt derzeit in einem separaten freigegebenen KI-Chat. SharePoint wird nicht direkt angebunden; der Partnerordner wird als ZIP verwendet. Die finale gestalterische und rechtliche Prüfung bleibt manuell.",
        fill=WARN,
    )
    add_para(
        doc,
        "Eine spätere API-Anbindung kann den separaten KI-Schritt auf einen Button reduzieren. Dafür werden ein zentral verwalteter API-Zugang, Datenschutzfreigabe, Kostenregelung, Fehlerbehandlung und Abnahmetests mit mehreren Partnern benötigt. Die jetzige JSON- und Promptstruktur kann dabei weiterverwendet werden.",
    )

    doc.add_heading("9. Dateien pro Partner", level=1)
    add_bullet(doc, "Partnerordner bzw. Partner-ZIP")
    add_bullet(doc, "KI-Textauftrag als TXT")
    add_bullet(doc, "KI-Ergebnis als Text-JSON")
    add_bullet(doc, "Finale Partner-JSON aus dem Mailing-Assistenten")
    add_bullet(doc, "Befüllte und bearbeitbare InDesign-Datei")
    add_bullet(doc, "Befüllreport, Prüf-PDF und finale Druckdaten")

    doc.add_heading("10. Definition of Done", level=1)
    add_para(
        doc,
        "Ein Partnerlauf ist abgeschlossen, wenn alle vier Varianten mit Vorder- und Rückseite befüllt, fachlich finalisiert, intern geprüft und als druckfähige Daten bereitgestellt sind. Die Zeitmessung muss alle vereinbarten Arbeitsschritte eindeutig ausweisen.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
