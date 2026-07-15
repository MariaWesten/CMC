from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


OUT = Path("outputs/Leistungsbeschreibung_URL_QR_Code_Tool.docx")

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F3F4F6"
BORDER = "D9E2F3"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, bold=False, color=None, size=None, italic=False, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, color=BLUE, bold=True)
    return p


def add_note(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ["top", "left", "bottom", "right"]:
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), "C7D2E3")
        borders.append(node)
    p_pr.append(borders)
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=DARK)
    r2 = p.add_run(text)
    set_run_font(r2, color=DARK)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].width = Inches(widths[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(hdr[i], LIGHT_BLUE)
        set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.10
            r = p.add_run(value)
            set_run_font(r, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def set_document_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def main():
    doc = Document()
    set_document_styles(doc)

    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Leistungsbeschreibung URL-/QR-Code-Tool"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "CMC | Stand 23.06.2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Leistungsbeschreibung URL-/QR-Code-Tool")
    set_run_font(r, size=22, color=DARK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("Arbeitsstand zur Abstimmung des erwarteten Funktionsumfangs und als Grundlage für eine Anfrage an Short.io")
    set_run_font(r, size=12.5, color=MUTED)

    meta = [
        ("An", "Robert"),
        ("Von", "Maria"),
        ("Stand", "23.06.2026"),
        ("Ziel", "Gemeinsames Verständnis des gewünschten URL-/QR-Code-Tools herstellen und anschließend Short.io mit dieser Leistungsbeschreibung anfragen."),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, bold=True, color=DARK)
        r2 = p.add_run(value)
        set_run_font(r2, color=DARK)

    add_note(
        doc,
        "Kurzfazit",
        "Die bisherige Anforderungsmatrix ist die Grundlage. Für die Anfrage an Short.io sollte sie jedoch als klare Leistungsbeschreibung formuliert werden: Welche Funktionen brauchen wir, welche Mengen müssen getragen werden und welche vertraglichen/betrieblichen Punkte müssen bestätigt werden.",
    )

    add_heading(doc, "1. Zielbild des Tools", 1)
    add_para(
        doc,
        "Das URL-/QR-Code-Tool soll Malin langfristig als zentrale Lösung für die kampagnenweise Erstellung, Verwaltung und Auswertung individueller Tracking-Links ablösen. Der Schwerpunkt liegt nicht auf einem klassischen QR-Code-Design-Tool, sondern auf stabilen, massenhaft erzeugbaren Tracking-Links, die optional als QR-Code-Bild ausgegeben werden können.",
    )
    add_para(
        doc,
        "Das Tool soll den bestehenden CSM-/Data-/Produktionsprozess zuverlässig abbilden und gleichzeitig genug Struktur schaffen, um Fehler vor dem Versand von Mailings zu vermeiden.",
    )

    add_heading(doc, "2. Erwarteter Funktionsumfang", 1)
    add_table(
        doc,
        ["Bereich", "Funktionalität", "Einordnung"],
        [
            (
                "Kampagnen",
                "Anlage und Verwaltung von Kampagnen mit Kampagnenname, internem Identifier, Start-/Enddatum, verantwortlichem Bereich und Status.",
                "Muss-Funktion für Übersicht, Reporting und spätere Nachvollziehbarkeit.",
            ),
            (
                "CSV-Upload",
                "Upload einer Datei mit Gutschein-/Aktionscodes, Ziel-URLs und weiteren Kampagnenfeldern. Die konkrete Feldstruktur muss vorab abgestimmt werden.",
                "Muss-Funktion, weil der heutige Prozess stark auf CSV-Dateien basiert.",
            ),
            (
                "Validierung",
                "Prüfung auf Pflichtfelder, doppelte Codes, ungültige URLs, leere Ziel-URLs und offensichtliche Formatfehler vor der Linkgenerierung.",
                "Muss-Funktion zur Fehlervermeidung vor Produktion und Mailingversand.",
            ),
            (
                "Tracking-Links",
                "Generierung individueller, stabiler Tracking-Links je Datensatz über eine definierte Tracking-Domain.",
                "Kernfunktion des Tools.",
            ),
            (
                "Redirect",
                "Zuverlässige Weiterleitung vom Tracking-Link zur hinterlegten Ziel-URL ohne Werbung, Delay oder externe Zwischenseite.",
                "Muss-Funktion, da fehlerhafte Links Mailings direkt gefährden.",
            ),
            (
                "QR-Code-Ausgabe",
                "Optional je Tracking-Link ein QR-Code-Bild erzeugen oder die QR-Code-Erstellung über Exportdaten ermöglichen.",
                "Wichtig für Produktion, aber fachlich nachgelagert gegenüber stabilen Tracking-Links.",
            ),
            (
                "Export",
                "Export der erzeugten Tracking-Links, Originaldaten, Ziel-URLs, Codes, Kampagnenfelder und optional QR-Dateien/QR-Referenzen.",
                "Muss-Funktion für CSM, Data und Produktion.",
            ),
            (
                "Reporting",
                "Klick-/Scan-Auswertung pro Kampagne, Link, Code und Zeitraum; zusätzlich exportierbare Roh- oder Aggregatdaten.",
                "Muss-Funktion, Detailtiefe mit Data/CSM zu klären.",
            ),
        ],
        [1.25, 3.45, 1.80],
    )

    add_heading(doc, "3. Mindestanforderungen an Betrieb und Zuverlässigkeit", 1)
    add_bullet(doc, "Die Links müssen auch bei großen Kampagnen stabil erreichbar sein.")
    add_bullet(doc, "Linkgenerierung und Export dürfen den Mailingprozess nicht blockieren; bei Fehlern braucht es klare Fehlermeldungen und Wiederholbarkeit.")
    add_bullet(doc, "Monitoring/Uptime-Prüfung ist erforderlich, weil Linkausfälle direkt auf laufende Mailings wirken.")
    add_bullet(doc, "Es braucht einen definierten Fallback- und Supportprozess: Wer reagiert bei Ausfall, wie schnell und über welchen Kanal?")
    add_bullet(doc, "Exports und Reporting müssen reproduzierbar sein, damit Ergebnisse später nachvollzogen werden können.")

    add_heading(doc, "4. Mengenmodell für Anbieteranfrage und Kostenprüfung", 1)
    add_para(
        doc,
        "Für die Bewertung von Short.io und anderen Markttools muss klar sein, für welches Mengengerüst das Angebot gelten soll. Die bisherigen Annahmen sind ein Arbeitsstand und sollten von Robert/CMC bestätigt oder angepasst werden.",
    )
    add_table(
        doc,
        ["Parameter", "Arbeitsannahme / zu klären", "Warum relevant"],
        [
            ("Kampagnen pro Monat", "Aktuell ca. 20 Kampagnen/Monat; bitte bestätigen.", "Basis für Planung, Betrieb und Support."),
            ("Links pro Kampagne", "Ca. 10.000 bis 100.000 individuelle Links je Kampagne.", "Haupttreiber für Tool-Limits und Anbieterpreise."),
            ("Neue Links pro Monat", "Ca. 200.000 bis 2.000.000 neue Links/Monat.", "Entscheidend für Enterprise-Preise und technische Skalierung."),
            ("Klick-/Scanvolumen", "Noch offen: erwartete Klicks/Scans pro Kampagne und Monat.", "Viele Anbieter limitieren nicht nur Links, sondern auch Events/Klicks."),
            ("Link-Laufzeit", "Zu klären: Wie lange müssen Links aktiv bleiben und reportbar sein?", "Relevant für Datenhaltung, Archivierung und Kosten."),
            ("Mittelfristiger Stand", "Zu klären: erwartetes Wachstum in 6-12 Monaten.", "Verhindert eine zu kleine oder schnell zu teure Lösung."),
        ],
        [1.55, 2.55, 2.40],
    )

    add_heading(doc, "5. Datenschutz, AVV/DPA und Datenumfang", 1)
    add_para(
        doc,
        "Es werden voraussichtlich keine sensiblen personenbezogenen Daten verarbeitet. Trotzdem können bei Tracking-Links technische Nutzungsdaten entstehen, z. B. IP-Adresse, Zeitpunkt, User-Agent/Gerät, Kampagne, Code, Ziel-URL und Server-Logs. Deshalb sollte kein großer Datenschutzprozess, aber ein kurzer Datenschutz-/AVV-Check eingeplant werden.",
    )
    add_bullet(doc, "AVV = Auftragsverarbeitungsvertrag; DPA = Data Processing Agreement.")
    add_bullet(doc, "Zu prüfen: Serverstandort, gespeicherte Trackingdaten, Speicherfristen, Exportrechte, Datenlöschung und Datenminimierung.")
    add_bullet(doc, "Bei Eigenbau kann stärker gesteuert werden, welche Daten überhaupt gespeichert werden.")

    add_heading(doc, "6. Mögliche spätere Erweiterungen beim Eigenbau", 1)
    add_para(
        doc,
        "Wenn CMC sich für einen Eigenbau entscheidet, sollte die erste Version bewusst schlank bleiben. Gleichzeitig kann die Architektur so angelegt werden, dass spätere Erweiterungen möglich sind, ohne die Stabilität des Kernprozesses zu gefährden.",
    )
    add_bullet(doc, "Rollen und Rechte für CSM, Data, Produktion und Administration.")
    add_bullet(doc, "Freigabe-Workflow vor finalem Export oder Produktionsfreigabe.")
    add_bullet(doc, "Kampagnen-Dashboard mit Status, Linkanzahl, Fehlern, Klick-/Scanentwicklung und Exporthistorie.")
    add_bullet(doc, "Automatische Plausibilitätschecks für Ziel-URLs, Dubletten, fehlende Felder und ungewöhnliche Werte.")
    add_bullet(doc, "Integrationen zu internen Daten-/CRM-/Produktionsprozessen, sofern später sinnvoll.")
    add_bullet(doc, "Langfristig wäre sogar eine Produktisierung/Abo-Variante theoretisch denkbar. Das sollte aber nicht Haupttreiber der aktuellen Entscheidung sein, sondern separat geprüft werden: Business Case, Mandantenfähigkeit, Support, SLA, Datenschutz, Abrechnung und Produktverantwortung.")

    add_heading(doc, "7. Fragen an Short.io", 1)
    add_table(
        doc,
        ["Thema", "Frage an Short.io"],
        [
            ("Volumen", "Können 200.000 bis 2.000.000 neue Tracking-Links pro Monat zuverlässig erstellt und betrieben werden? Falls ja: in welchem Plan und zu welchen Kosten?"),
            ("Bulk/API", "Welche Möglichkeiten gibt es für CSV-Upload, Bulk-Erstellung oder API-basierte Linkgenerierung? Gibt es Limits pro Datei, API-Rate-Limits oder Automatisierungslimits?"),
            ("Redirects", "Sind Redirects dauerhaft stabil, ohne Werbung, Delay oder Zwischenseite? Gibt es ein SLA und wie wird Verfügbarkeit garantiert?"),
            ("QR-Codes", "Können QR-Code-Bilder je Link erzeugt oder exportiert werden? Gibt es Mengen- oder Downloadlimits?"),
            ("Reporting", "Welche Klick-/Scan-Daten sind verfügbar? Können Reports pro Kampagne, Link und Zeitraum exportiert werden? Gibt es Eventlimits oder Aufbewahrungsfristen?"),
            ("Datenexport", "Welche Exportformate gibt es für Links, Kampagnen, QR-Daten und Reporting? Was passiert bei Kündigung mit Daten und Links?"),
            ("Domain", "Können eigene Tracking-Domains genutzt werden? Wie viele Domains sind enthalten und wie wird DNS eingerichtet?"),
            ("Datenschutz", "Gibt es AVV/DPA, EU-Regionen oder Angaben zu Serverstandort, Logdaten und Speicherfristen?"),
            ("Support/Betrieb", "Welche Supportlevel, Reaktionszeiten, Monitoringoptionen und Eskalationswege gibt es?"),
            ("Preis", "Bitte Angebot für heutigen Stand und mittelfristiges Zielvolumen erstellen, inklusive aller Zusatzkosten für Links, Klicks/Events, Domains, Nutzer, API, Exporte, SLA und Support."),
        ],
        [1.45, 5.05],
    )

    add_heading(doc, "8. Punkte, die intern vor der Anfrage finalisiert werden sollten", 1)
    add_bullet(doc, "Finales CSV-Format: Pflichtfelder, optionale Felder, Code-/URL-Logik und erwartete Exportspalten.")
    add_bullet(doc, "Heutiges und mittelfristiges Mengengerüst: Kampagnen, Links, Klicks/Scans, Laufzeit der Links.")
    add_bullet(doc, "Reporting-Erwartung: Welche Auswertungen braucht CSM/Data wirklich?")
    add_bullet(doc, "Betriebsmodell beim Eigenbau: intern/managed oder externer Betrieb, z. B. Rent your Admin.")
    add_bullet(doc, "Kostenblock Rent your Admin konkret anfragen: Setup, Monitoring, Backups, Incident-Reaktion, Reaktionszeiten, monatliche Pauschale und Zusatzkosten.")

    add_heading(doc, "9. Vorschlag für das weitere Vorgehen", 1)
    add_bullet(doc, "Robert prüft, ob der oben beschriebene Funktionsumfang dem erwarteten Zielbild entspricht.")
    add_bullet(doc, "Nach Freigabe wird Short.io mit dieser Leistungsbeschreibung angefragt.")
    add_bullet(doc, "Parallel werden die offenen Mengen- und Betriebsparameter intern geklärt.")
    add_bullet(doc, "Nach Rückmeldung von Short.io werden Kosten, Funktionsabdeckung, Risiken und Eigenbau-Alternative aktualisiert gegenübergestellt.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
